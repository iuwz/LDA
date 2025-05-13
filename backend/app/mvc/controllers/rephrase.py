# backend/app/mvc/controllers/rephrase.py

import logging
import json
import difflib
import os
from fastapi import HTTPException
from motor.motor_asyncio import AsyncIOMotorDatabase
from io import BytesIO
from typing import Optional, Dict, Any
from bson import ObjectId
from datetime import datetime

from backend.app.mvc.controllers.documents import (
    get_document_record,
    open_gridfs_file,
    store_document_record,
    upload_file_to_gridfs,
)
from backend.app.core.openai_client import call_gpt

# Try optional imports
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

try:
    from PyPDF2 import PdfReader
except ImportError:
    PdfReader = None

logger = logging.getLogger(__name__)


def compute_changes(original: str, revised: str) -> list[Dict[str, str]]:
    """Compute word-level 'replace' diffs between two texts."""
    orig_tokens = original.split()
    rev_tokens = revised.split()
    sm = difflib.SequenceMatcher(None, orig_tokens, rev_tokens)
    changes = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "replace":
            changes.append({
                "original": " ".join(orig_tokens[i1:i2]),
                "revised":  " ".join(rev_tokens[j1:j2]),
            })
    return changes


async def extract_full_text_from_stream(stream, filename: str) -> str:
    content = await stream.read()
    ext = filename.rsplit(".", 1)[-1].lower()

    # DOCX
    if ext == "docx" and DocxDocument:
        doc = DocxDocument(BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs)

    # PDF
    if ext == "pdf" and PdfReader:
        reader = PdfReader(BytesIO(content))
        if reader.is_encrypted:
            raise HTTPException(status_code=422, detail="Encrypted PDF")
        return "".join(page.extract_text() or "" for page in reader.pages)

    # fallback
    try:
        return content.decode("utf-8")
    except:
        return content.decode("latin-1", errors="ignore")


def create_simple_docx_from_text(text: str) -> bytes:
    if not DocxDocument:
        raise HTTPException(status_code=500, detail="python-docx missing")
    doc = DocxDocument()
    for line in text.split("\n"):
        doc.add_paragraph(line)
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


async def run_rephrase_tool(
    db: AsyncIOMotorDatabase,
    user_id: str,
    style: str,
    document_text: Optional[str] = None,
    doc_id: Optional[str] = None,
) -> Dict[str, Any]:

    if (document_text is None) == (doc_id is None):
        raise HTTPException(status_code=400, detail="Provide either text or doc_id")

    # ── Build the base report record with a proper created_at ─────────
    report_record: Dict[str, Any] = {
        "user_id": user_id,
        "style":   style,
        "created_at": datetime.utcnow(),
    }

    # ── Document Mode ────────────────────────────────────────────────
    if doc_id:
        rec = await get_document_record(db, doc_id)
        grid_out, orig_fn = await open_gridfs_file(db, rec["file_id"])
        original = await extract_full_text_from_stream(grid_out, orig_fn)

        # AI call
        system_msg = (
            f"You are a legal rewriting AI. Rephrase the following "
            f"to a {style} style, preserving meaning. Return only plain text."
        )
        revised = call_gpt(
            prompt=original, system_message=system_msg, temperature=0.4
        ).strip() or original

        # build and store new .docx
        new_bytes = create_simple_docx_from_text(revised)
        base, _ = os.path.splitext(orig_fn)
        new_fn = f"{base}_rephrased.docx"

        file_id    = await upload_file_to_gridfs(db, new_bytes, new_fn)
        new_doc_id = await store_document_record(db, user_id, new_fn, file_id)

        changes = compute_changes(original, revised)

        report_record.update({
            "original_content_info":        orig_fn,
            "rephrased_output_summary":     new_fn,
            "original_doc_id":              doc_id,
            "rephrased_doc_id":             new_doc_id,
            "changes":                      changes,
        })

        res = await db.rephrase_reports.insert_one(report_record)
        rid = str(res.inserted_id)

        return {
            "report_id":           rid,
            "rephrased_doc_id":    new_doc_id,
            "rephrased_doc_filename": new_fn,
            "changes":             changes,
        }

    # ── Text Mode ───────────────────────────────────────────────────
    original = document_text or ""
    system_msg = (
    f"You are an advanced Saudi legal-drafting AI. Rewrite the given text in the requested "
    f"{style} style—whether formal Arabic legal prose, plain-language English, or any other—"
    f"without altering the substance, legal effect, citations, or cross-references. "
    f"Do not add new terms or remove existing obligations. "
    f"Return **valid JSON only** (no markdown, no commentary) in exactly this structure:\n"
    f'{{"rephrased_text":"..."}}'
)

    ai = call_gpt(prompt=original, system_message=system_msg, temperature=0.4)
    try:
        parsed = json.loads(ai)
        revised = parsed.get("rephrased_text", original).strip()
    except:
        revised = ai.strip() or original

    changes = compute_changes(original, revised)

    report_record.update({
        "original_content_info":    original,
        "rephrased_output_summary": revised,
        "original_doc_id":          None,
        "rephrased_doc_id":         None,
        "changes":                  changes,
    })

    res = await db.rephrase_reports.insert_one(report_record)
    rid = str(res.inserted_id)

    return {
        "report_id":   rid,
        "rephrased_text": revised,
        "changes":     changes,
    }

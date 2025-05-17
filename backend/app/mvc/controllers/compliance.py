# backend/app/mvc/controllers/compliance.py
from __future__ import annotations

import datetime as _dt
import json
import logging
import re
from io import BytesIO
from typing import Any, Dict, List, Optional

from bson import ObjectId
from fastapi import HTTPException
from fastapi.concurrency import run_in_threadpool
from motor.motor_asyncio import AsyncIOMotorDatabase

from docx import Document
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Table,
    TableStyle,
    Spacer,
)
from reportlab.lib.styles import getSampleStyleSheet

from backend.app.core.openai_client import call_gpt
from backend.app.mvc.controllers.documents import (
    get_document_record,
    open_gridfs_file,
    extract_full_text_from_stream,
    upload_file_to_gridfs,
    store_document_record,
)
from backend.app.mvc.models.compliance import ComplianceIssue

logger = logging.getLogger(__name__)

# size guard: 12 000 characters ≈ 3 k tokens → stays well below GPT-3.5’s 4 k ctx
CHUNK_CHAR_LIMIT = 12_000


# ═════════════════════ helper utilities ══════════════════════
def _split_into_chunks(text: str, size: int = CHUNK_CHAR_LIMIT) -> List[str]:
    """
    Greedy splitter that *tries* not to break sentences.

    We split on line boundaries; this keeps clause structure intact and
    avoids chopping right in the middle of a sentence or word.
    """
    out: list[str] = []
    buf: list[str] = []
    cur_len = 0
    for line in text.splitlines(keepends=True):
        buf.append(line)
        cur_len += len(line)
        if cur_len >= size:
            out.append("".join(buf))
            buf, cur_len = [], 0
    if buf:
        out.append("".join(buf))
    return out


def _deduplicate_issues(raw: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """
    Remove duplicates that may appear when several chunks flag the same clause.

    Key: (rule_id + first 50 chars of snippet) – good enough for our purpose.
    """
    unique: list[Dict[str, Any]] = []
    seen: set[str] = set()
    for item in raw:
        key = f"{item.get('rule_id')}|{(item.get('extracted_text_snippet') or '')[:50]}"
        if key not in seen:
            unique.append(item)
            seen.add(key)
    return unique


def _extract_snippet(text: str, pattern: str, win: int = 100) -> str:
    """
    Return an excerpt around the first occurrence of *pattern*.

    If the pattern is not found → empty string.
    """
    m = re.search(pattern, text, re.IGNORECASE)
    if not m:
        return ""
    s, e = max(0, m.start() - win), min(len(text), m.end() + win)
    return ("…" if s else "") + text[s:e] + ("…" if e < len(text) else "")


def _heuristic_score(issues: List[Dict[str, Any]]) -> int:
    """
    Very simple fallback scoring:
        – start at 100
        – each “Issue Found” deducts 20
        – never below 0
    """
    bad = sum(1 for i in issues if i.get("status", "").lower() == "issue found")
    return max(0, 100 - bad * 20)


async def _store_pdf(
    db: AsyncIOMotorDatabase,
    user_id: str,
    buf: BytesIO,
    filename: str,
) -> str:
    """
    Upload the in-memory PDF buffer to GridFS, create a document record,
    return the newly created doc-id (string).
    """
    buf.seek(0)
    grid_id = await upload_file_to_gridfs(db, buf.read(), filename)
    return await store_document_record(db, user_id, filename, grid_id)


# ═════════════════════ main entry point ══════════════════════
async def run_compliance_check(
    db: AsyncIOMotorDatabase,
    user_id: str,
    *,
    document_text: Optional[str] = None,
    doc_id: Optional[str] = None,
) -> Dict[str, Any]:
    """
    Analyse the given document for Saudi PDPL / contract-law compliance.

    You can either pass:
        • raw text   – via `document_text`
        • an uploaded document id – via `doc_id`
    Exactly one must be provided.

    The routine:
        1. Ensures we have the plain text (extracts from GridFS if needed)
        2. Splits it into chunks so each GPT call stays within context limits
        3. Calls GPT for each chunk (in a background thread)
        4. Deduplicates / normalises the issues
        5. Stores the report in MongoDB
        6. Generates & uploads a PDF report
        7. Returns the consolidated result payload
    """
    # ────────────────── sanity checks ──────────────────
    if not (document_text or doc_id):
        raise HTTPException(400, "Either document_text or doc_id must be provided.")
    if document_text and doc_id:
        raise HTTPException(400, "Provide only one of document_text or doc_id.")

    # ────────────────── load text if only doc-id was provided ──────────────────
    if doc_id:
        rec = await get_document_record(db, doc_id)
        stream, fname = await open_gridfs_file(db, rec["file_id"])
        document_text = await extract_full_text_from_stream(stream, fname)
        if document_text.startswith("Error"):
            raise HTTPException(422, document_text)

    # ────────────────── GPT prompt (system message) ──────────────────
    system_message = """
You are an expert Saudi-law compliance AI assistant. Review the contract text
line-by-line and flag *only* those clauses that breach either:

  • the KSA Personal Data Protection Law (PDPL, Royal Decree M/19 of 9 Safar 1443 H)
  • mandatory Saudi contract-law principles derived from Sharia, Royal Decrees,
    the Civil Transactions Law, or any binding ministerial regulations.

For every non-compliant clause, output EXACTLY these keys—nothing more:

  • rule_id                – concise, uppercase identifier with no spaces
  • description            – brief explanation of the violation
  • status                 – always the string "Issue Found"
  • extracted_text_snippet – verbatim offending language from the contract

If the contract contains zero violations, return ONLY this exact JSON:
{"issues":[]}

Respond with VALID JSON only—no markdown and no additional commentary.
""".strip()

    # ────────────────── call GPT chunk-by-chunk ──────────────────
    raw_issues: list[Dict[str, Any]] = []

    for idx, chunk in enumerate(_split_into_chunks(document_text)):
        try:
            resp = await call_gpt(
                chunk,
                system_message=system_message,
                model="o4-mini",
                temperature=0.0,
                max_tokens=16384,
            )
            parsed: Any = json.loads(resp) if resp else {}
            if isinstance(parsed, dict):
                raw_issues.extend(parsed.get("issues", []))
            elif isinstance(parsed, list):
                raw_issues.extend(parsed)
        except json.JSONDecodeError:
            logger.warning("Chunk %d → GPT returned invalid JSON: %r", idx + 1, resp)
        except Exception:
            logger.exception("Chunk %d → GPT call failed", idx + 1)

    # ────────────────── deduplicate & normalise ──────────────────
    raw_issues = _deduplicate_issues(raw_issues)

    issues: List[Dict[str, Any]] = []
    for raw in raw_issues:
        if not isinstance(raw, dict):
            continue
        # ensure minimal required keys
        raw["status"] = "Issue Found"
        if not raw.get("extracted_text_snippet"):
            pat = raw.get("rule_id", "")
            raw["extracted_text_snippet"] = _extract_snippet(document_text, pat) or None
        try:
            issues.append(ComplianceIssue(**raw).model_dump())
        except Exception:
            # validation failed – skip silently
            continue

    # No issues → mark explicitly clean
    if not issues:
        issues = [
            {
                "rule_id": "NONE",
                "description": "No compliance issues detected.",
                "status": "OK",
                "extracted_text_snippet": None,
            }
        ]

    # ────────────────── compute score ──────────────────
    compliance_score = _heuristic_score(issues)

    # ────────────────── persist metadata ──────────────────
    preview = document_text[:500] + ("…" if len(document_text) > 500 else "")
    row = {
        "user_id": user_id,
        "document_text_preview": preview,
        "original_doc_id": doc_id,
        "issues": issues,
        "compliance_score": compliance_score,
        "timestamp": _dt.datetime.utcnow(),
    }
    ins = await db.compliance_reports.insert_one(row)
    report_id = str(ins.inserted_id)

    # ────────────────── generate PDF report ──────────────────
    pdf_buf: BytesIO = await run_in_threadpool(
        generate_compliance_report_pdf,
        {**row, "_id": report_id},
    )
    pdf_name = f"compliance_report_{report_id}.pdf"
    report_doc_id = await _store_pdf(db, user_id, pdf_buf, pdf_name)

    # add PDF metadata to DB record
    await db.compliance_reports.update_one(
        {"_id": ins.inserted_id},
        {"$set": {"report_doc_id": report_doc_id, "report_filename": pdf_name}},
    )

    # ────────────────── response payload ──────────────────
    return {
        "report_id": report_id,
        "issues": issues,
        "compliance_score": compliance_score,
        "report_doc_id": report_doc_id,
        "report_filename": pdf_name,
    }


# ═════════════════════ DB fetch / delete helpers ══════════════════════
async def get_compliance_report(
    db: AsyncIOMotorDatabase,
    report_id: str,
    user_id: str,
) -> Dict[str, Any]:
    """
    Retrieve a stored report, enforcing ownership.
    """
    if not ObjectId.is_valid(report_id):
        raise HTTPException(400, "Invalid report ID.")
    doc = await db.compliance_reports.find_one(
        {"_id": ObjectId(report_id), "user_id": user_id}
    )
    if not doc:
        raise HTTPException(404, "Report not found or access denied.")
    doc["_id"] = str(doc["_id"])
    return doc


# ═════════════════════ report generators (DOCX / PDF) ══════════════════════
def generate_compliance_report_docx(report_data: Dict[str, Any]) -> BytesIO:
    """
    Build a nicely formatted DOCX from the compliance data.
    """
    buf = BytesIO()
    document = Document()

    # Cover page
    document.add_heading("Compliance Report", level=0)
    meta = document.add_table(rows=5, cols=2, style="Light List Accent 1")
    meta.cell(0, 0).text = "Report ID:"
    meta.cell(0, 1).text = report_data.get("_id", "N/A")
    ts = report_data.get("timestamp")
    meta.cell(1, 0).text = "Generated At:"
    meta.cell(1, 1).text = ts.strftime("%Y-%m-%d %H:%M:%S UTC") if ts else "N/A"
    meta.cell(2, 0).text = "User:"
    meta.cell(2, 1).text = report_data.get("user_id", "N/A")
    meta.cell(3, 0).text = "Document:"
    meta.cell(3, 1).text = report_data.get("original_doc_id", "N/A")
    meta.cell(4, 0).text = "Compliance Score:"
    meta.cell(4, 1).text = str(report_data.get("compliance_score", "N/A"))
    document.add_page_break()

    # Summary of Findings
    document.add_heading("1. Summary of Findings", level=1)
    counts = {"OK": 0, "Issue Found": 0}
    for issue in report_data["issues"]:
        counts[issue["status"]] = counts.get(issue["status"], 0) + 1

    tbl = document.add_table(rows=2, cols=2, style="Light Grid Accent 2")
    tbl.rows[0].cells[0].text = "Issues Found"
    tbl.rows[1].cells[0].text = str(counts["Issue Found"])
    tbl.rows[0].cells[1].text = "Clean"
    tbl.rows[1].cells[1].text = str(counts["OK"])
    document.add_page_break()

    # Detailed Issues
    document.add_heading("2. Detailed Issues", level=1)
    for idx, issue in enumerate(report_data["issues"], start=1):
        document.add_heading(f"Issue {idx}: {issue['rule_id']}", level=2)
        t = document.add_table(rows=4, cols=2, style="Light List Accent 1")
        t.cell(0, 0).text = "Status"
        t.cell(0, 1).text = issue["status"]
        t.cell(1, 0).text = "Description"
        t.cell(1, 1).text = issue["description"]
        t.cell(2, 0).text = "Snippet"
        t.cell(2, 1).text = issue.get("extracted_text_snippet") or "–"
        t.cell(3, 0).text = "Rule ID"
        t.cell(3, 1).text = issue["rule_id"]
        document.add_paragraph("")

    # footer
    footer = document.sections[-1].footer
    footer.paragraphs[0].text = (
        f"Generated on {_dt.datetime.utcnow():%Y-%m-%d %H:%M:%S UTC}"
    )
    footer.paragraphs[0].alignment = 1  # centered

    document.save(buf)
    buf.seek(0)
    return buf


def generate_compliance_report_pdf(report_data: Dict[str, Any]) -> BytesIO:
    """
    Build a PDF report using ReportLab.
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=letter)
    styles = getSampleStyleSheet()
    elems: List[Any] = []

    # Title + meta
    elems.append(Paragraph("Compliance Report", styles["Title"]))
    elems.append(Spacer(1, 12))
    ts = report_data.get("timestamp")
    meta_tbl = Table(
        [
            ["Report ID:", report_data.get("_id", "N/A")],
            ["Generated At:", ts.strftime("%Y-%m-%d %H:%M:%S UTC") if ts else "N/A"],
            ["User:", report_data.get("user_id", "N/A")],
            ["Compliance Score:", str(report_data.get("compliance_score", "N/A"))],
        ],
        colWidths=[120, 330],
    )
    meta_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]
        )
    )
    elems.append(meta_tbl)
    elems.append(Spacer(1, 24))

    # Summary
    elems.append(Paragraph("1. Summary of Findings", styles["Heading2"]))
    counts = {"OK": 0, "Issue Found": 0}
    for issue in report_data["issues"]:
        counts[issue["status"]] += 1

    sum_tbl = Table(
        [
            ["Status", "Count"],
            ["Issue Found", str(counts["Issue Found"])],
            ["OK", str(counts["OK"])],
        ],
        colWidths=[120, 120],
    )
    sum_tbl.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.lightblue),
                ("BOX", (0, 0), (-1, -1), 1, colors.black),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ]
        )
    )
    elems.append(sum_tbl)
    elems.append(Spacer(1, 24))

    # Detailed Issues
    elems.append(Paragraph("2. Detailed Issues", styles["Heading2"]))
    for idx, issue in enumerate(report_data["issues"], start=1):
        elems.append(Paragraph(f"Issue {idx}: {issue['rule_id']}", styles["Heading3"]))
        det_tbl = Table(
            [
                ["Status", Paragraph(issue["status"], styles["Normal"])],
                ["Description", Paragraph(issue["description"], styles["Normal"])],
                [
                    "Snippet",
                    Paragraph(issue.get("extracted_text_snippet") or "–", styles["Normal"]),
                ],
                ["Rule ID", Paragraph(issue["rule_id"], styles["Normal"])],
            ],
            colWidths=[120, 330],
        )
        det_tbl.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.whitesmoke),
                    ("BOX", (0, 0), (-1, -1), 1, colors.black),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ]
            )
        )
        elems.append(det_tbl)
        elems.append(Spacer(1, 12))

    elems.append(Spacer(1, 24))
    elems.append(
        Paragraph(
            f"Generated on {_dt.datetime.utcnow():%Y-%m-%d %H:%M:%S UTC}",
            styles["Normal"],
        )
    )

    doc.build(elems)
    buf.seek(0)
    return buf
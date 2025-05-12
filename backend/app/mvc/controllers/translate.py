# backend/app/mvc/controllers/translate.py
"""
Translation controller.

Adds persistent storage + GridFS upload so that:
• text translations are kept (type='text')
• file translations save the generated DOCX in GridFS (type='doc')
  and patch the translation_reports row with result_doc_id +
  translated_filename.
"""

from __future__ import annotations

import logging, re, datetime as _dt
from io import BytesIO
from typing import Dict, Any, Optional, Tuple

from bson import ObjectId
from fastapi import HTTPException, UploadFile
from fastapi.concurrency import run_in_threadpool
from motor.motor_asyncio import AsyncIOMotorDatabase
from dotenv import load_dotenv
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches

from backend.app.core.openai_client import call_gpt
from backend.app.mvc.controllers.rephrase import extract_full_text_from_stream
from backend.app.mvc.controllers.documents import (
    upload_file_to_gridfs,
    store_document_record,
)

load_dotenv()
logger = logging.getLogger(__name__)

# ───────────────────────── INTERNAL ─────────────────────────
async def _upload_docx(
    db: AsyncIOMotorDatabase,
    user_id: str,
    buffer: BytesIO,
    filename: str,
) -> str:
    """Upload *buffer* to GridFS and return the document’s doc-id."""
    buffer.seek(0)
    grid_id = await upload_file_to_gridfs(db, buffer.read(), filename)
    return await store_document_record(db, user_id, filename, grid_id)


async def _insert_base_row(
    db: AsyncIOMotorDatabase,
    user_id: str,
    original_text: str,
    target_lang: str,
    *,
    original_doc_id: Optional[str] = None,
) -> str:
    row = {
        "user_id": user_id,
        "type": "text" if original_doc_id is None else "doc",
        "target_lang": target_lang,
        "original_text": (original_text[:1000] + "…")
        if len(original_text) > 1000
        else original_text,
        "original_doc_id": original_doc_id,
        "timestamp": _dt.datetime.utcnow(),
    }
    res = await db.translation_reports.insert_one(row)
    return str(res.inserted_id)


# ─────────────────────────  PUBLIC  ──────────────────────────
async def run_translation_tool(
    db: AsyncIOMotorDatabase,
    document_text: str,
    target_lang: str,
    user_id: str,
) -> Dict[str, Any]:
    """
    Translate raw *document_text* into *target_lang* and store the row
    (type='text').
    """
    if not document_text:
        raise HTTPException(400, "Document text is required.")
    if not target_lang:
        raise HTTPException(400, "target_lang is required.")

    logger.info(
        "Translating %d chars for user=%s → %s",
        len(document_text),
        user_id,
        target_lang,
    )

    report_id = await _insert_base_row(
        db, user_id, document_text, target_lang, original_doc_id=None
    )

    sys_msg = (
        f"You are a professional legal translator. "
        f"Translate into {target_lang.upper()}."
    )
    user_msg = (
        f"Translate this legal document into {target_lang.upper()}:\n\n"
        f"{document_text}"
    )

    try:
        translated_text = call_gpt(
            prompt=user_msg,
            system_message=sys_msg,
        )
        if not translated_text:
            raise Exception("Translation response was empty")

        await db.translation_reports.update_one(
            {"_id": ObjectId(report_id)},
            {"$set": {"translated_text": translated_text}},
        )

    except Exception as e:
        logger.error("Translation failed: %s", e, exc_info=True)
        raise HTTPException(500, "Translation failed due to an internal error")

    return {"report_id": report_id, "translated_text": translated_text}


async def run_file_translation_tool(
    db: AsyncIOMotorDatabase,
    file: UploadFile,
    target_lang: str,
    user_id: str,
) -> Tuple[bytes, str, str]:
    """
    Translate an uploaded file; returns (blob_bytes, translated_filename,
    report_id) so the route can stream the DOCX and the front-end can display.
    """
    raw = await file.read()

    class _AsyncBuf:
        def __init__(self, b: bytes):
            self._b = b

        async def read(self):
            return self._b

    extracted_text = await extract_full_text_from_stream(
        _AsyncBuf(raw), file.filename
    )
    if extracted_text.startswith("Error:"):
        raise HTTPException(422, extracted_text)

    report_id = await _insert_base_row(
        db,
        user_id,
        extracted_text,
        target_lang,
        original_doc_id=None,  # could also save original upload
    )

    sys_msg = (
        f"You are a professional legal translator. "
        f"Translate into {target_lang.upper()}."
    )
    user_msg = (
        f"Translate this legal document into {target_lang.upper()}:\n\n"
        f"{extracted_text}"
    )

    try:
        translated_text = call_gpt(
            prompt=user_msg,
            system_message=sys_msg,
        )
        if not translated_text:
            raise Exception("Translation response was empty")
    except Exception as e:
        logger.error("File translation failed: %s", e, exc_info=True)
        raise HTTPException(500, "Internal translation error")

    pdf_buf: BytesIO = await run_in_threadpool(
        _build_docx, translated_text, target_lang
    )

    translated_filename = (
        f"translated_{file.filename.rsplit('.', 1)[0]}_"
        f"{target_lang.lower()}.docx"
    )
    doc_id = await _upload_docx(db, user_id, pdf_buf, translated_filename)

    await db.translation_reports.update_one(
        {"_id": ObjectId(report_id)},
        {
            "$set": {
                "translated_text": translated_text,
                "result_doc_id": doc_id,
                "translated_filename": translated_filename,
                "type": "doc",  # ← FIXED: mark this row as a document translation
            }
        },
    )

    logger.info(
        "Stored translated DOCX %s (doc_id=%s)", translated_filename, doc_id
    )
    return pdf_buf.getvalue(), translated_filename, report_id


# ─────────────────────  Helper: DOCX builder  ─────────────────────
def _build_docx(translated_content: str, target_lang: str) -> BytesIO:
    """Return DOCX buffer (report-style layout, RTL aware)."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Inches(1)
        section.left_margin = section.right_margin = Inches(1)

    header = doc.add_heading(
        f"Translated Document ({target_lang.upper()})", level=0
    )
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    header.runs[0].font.name = "Arial"
    header.runs[0].font.size = Pt(16)

    doc.add_paragraph().add_run().add_break()

    for line in translated_content.splitlines():
        raw = line.strip()
        if not raw:
            continue
        clean = raw.replace("**", "").strip()

        match = re.match(r"^(\d+)\.\s*(.+)", clean)
        if match:
            num, heading = match.groups()
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            r1 = p.add_run(f"{heading} ")
            r1.bold = True
            r1.font.name = "Arial"
            r1.font.size = Pt(14)
            r2 = p.add_run(f"{num}.")
            r2.bold = True
            r2.font.name = "Arial"
            r2.font.size = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        else:
            p = doc.add_paragraph(clean)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf
